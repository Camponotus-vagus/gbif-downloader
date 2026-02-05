"""
Output generators for various file formats.

Supports: HTML, PDF, DOCX
"""

from pathlib import Path
from typing import Union, Optional
import tempfile
import webbrowser

from ..label_generator import LabelGenerator, LabelConfig


def generate_html(
    generator: LabelGenerator,
    output_path: Optional[Union[str, Path]] = None,
    open_in_browser: bool = False
) -> str:
    """Generate HTML file with labels.

    Args:
        generator: LabelGenerator with labels to render
        output_path: Path to save the HTML file. If None, returns HTML string only.
        open_in_browser: Whether to open the file in the default browser

    Returns:
        HTML content as string
    """
    config = generator.config
    html = _generate_html_content(generator)

    if output_path:
        path = Path(output_path)
        path.write_text(html, encoding="utf-8")

        if open_in_browser:
            webbrowser.open(f"file://{path.absolute()}")

    return html


def _generate_html_content(generator: LabelGenerator) -> str:
    """Generate the HTML content for labels."""
    config = generator.config

    # CSS styles
    css = f"""
    @page {{
        size: {config.page_width_mm}mm {config.page_height_mm}mm;
        margin: 0;
    }}

    * {{
        margin: 0;
        padding: 0;
        box-sizing: border-box;
    }}

    body {{
        font-family: {config.font_family}, sans-serif;
        font-size: {config.font_size_pt}pt;
        line-height: {config.line_spacing};
        -webkit-print-color-adjust: exact;
        print-color-adjust: exact;
    }}

    .page {{
        width: {config.page_width_mm}mm;
        height: {config.page_height_mm}mm;
        padding-top: {config.margin_top_mm}mm;
        padding-bottom: {config.margin_bottom_mm}mm;
        padding-left: {config.margin_left_mm}mm;
        padding-right: {config.margin_right_mm}mm;
        page-break-after: always;
        display: flex;
        flex-wrap: wrap;
        align-content: flex-start;
    }}

    .page:last-child {{
        page-break-after: auto;
    }}

    .label {{
        width: {config.label_width_mm}mm;
        height: {config.label_height_mm}mm;
        padding: 1mm;
        overflow: hidden;
        display: flex;
        flex-direction: column;
        justify-content: flex-start;
        border: 0.1mm solid #ddd;
    }}

    .label-content {{
        display: flex;
        flex-direction: column;
    }}

    .location-line {{
        font-size: {config.font_size_pt}pt;
        white-space: nowrap;
        overflow: hidden;
        text-overflow: ellipsis;
    }}

    .empty-line {{
        height: {config.font_size_pt * config.line_spacing}pt;
    }}

    .code {{
        font-size: {config.font_size_pt}pt;
        font-weight: normal;
    }}

    .date {{
        font-size: {config.font_size_pt}pt;
    }}

    .additional-info {{
        font-size: {config.font_size_pt * 0.9}pt;
        font-style: italic;
    }}

    /* Print styles */
    @media print {{
        body {{
            margin: 0;
            padding: 0;
        }}

        .page {{
            page-break-after: always;
            margin: 0;
        }}

        .label {{
            border: 0.1mm solid #ccc;
        }}

        .no-print {{
            display: none !important;
        }}
    }}

    /* Screen styles for preview */
    @media screen {{
        body {{
            background: #f0f0f0;
            padding: 20px;
        }}

        .page {{
            background: white;
            margin: 20px auto;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}

        .print-button {{
            position: fixed;
            top: 20px;
            right: 20px;
            padding: 10px 20px;
            background: #4CAF50;
            color: white;
            border: none;
            border-radius: 5px;
            cursor: pointer;
            font-size: 14px;
            z-index: 1000;
        }}

        .print-button:hover {{
            background: #45a049;
        }}

        .page-info {{
            text-align: center;
            margin: 10px;
            color: #666;
            font-size: 12px;
        }}
    }}
    """

    # Generate pages
    pages_html = []
    for page_num in range(generator.total_pages):
        grid = generator.get_labels_grid(page_num)
        labels_html = []

        for row in grid:
            for label in row:
                if label and not label.is_empty():
                    label_html = f"""
                    <div class="label">
                        <div class="label-content">
                            <div class="location-line">{_escape_html(label.location_line1)}</div>
                            <div class="location-line">{_escape_html(label.location_line2)}</div>
                            <div class="empty-line"></div>
                            <div class="code">{_escape_html(label.code)}</div>
                            <div class="date">{_escape_html(label.date)}</div>
                            {f'<div class="additional-info">{_escape_html(label.additional_info)}</div>' if label.additional_info else ''}
                        </div>
                    </div>
                    """
                else:
                    label_html = '<div class="label"></div>'
                labels_html.append(label_html)

        page_html = f"""
        <div class="page-info no-print">Pagina {page_num + 1} di {generator.total_pages}</div>
        <div class="page">
            {''.join(labels_html)}
        </div>
        """
        pages_html.append(page_html)

    html = f"""<!DOCTYPE html>
<html lang="it">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Etichette Entomologiche</title>
    <style>
        {css}
    </style>
</head>
<body>
    <button class="print-button no-print" onclick="window.print()">Stampa / Salva PDF</button>
    {''.join(pages_html)}
    <script>
        // Auto-adjust for print
        window.onbeforeprint = function() {{
            document.querySelectorAll('.no-print').forEach(el => el.style.display = 'none');
        }};
        window.onafterprint = function() {{
            document.querySelectorAll('.no-print').forEach(el => el.style.display = '');
        }};
    </script>
</body>
</html>
"""
    return html


def _escape_html(text: str) -> str:
    """Escape HTML special characters."""
    if not text:
        return ""
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&#39;")
    )


def generate_pdf(
    generator: LabelGenerator,
    output_path: Union[str, Path],
    open_after: bool = False
) -> Path:
    """Generate PDF file with labels.

    Requires weasyprint to be installed.

    Args:
        generator: LabelGenerator with labels to render
        output_path: Path to save the PDF file
        open_after: Whether to open the file after generation

    Returns:
        Path to the generated PDF file
    """
    try:
        from weasyprint import HTML, CSS
    except ImportError:
        raise ImportError(
            "weasyprint is required for PDF generation. "
            "Install with: pip install weasyprint\n"
            "Note: weasyprint requires additional system dependencies. "
            "See: https://doc.courtbouillon.org/weasyprint/stable/first_steps.html"
        )

    path = Path(output_path)
    html_content = _generate_html_content(generator)

    # Generate PDF
    html = HTML(string=html_content)
    html.write_pdf(path)

    if open_after:
        webbrowser.open(f"file://{path.absolute()}")

    return path


def generate_docx(
    generator: LabelGenerator,
    output_path: Union[str, Path],
    open_after: bool = False
) -> Path:
    """Generate Word document (.docx) with labels.

    Requires python-docx to be installed.

    Args:
        generator: LabelGenerator with labels to render
        output_path: Path to save the DOCX file
        open_after: Whether to open the file after generation

    Returns:
        Path to the generated DOCX file
    """
    try:
        from docx import Document
        from docx.shared import Mm, Pt
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX generation. "
            "Install with: pip install python-docx"
        )

    config = generator.config
    path = Path(output_path)

    doc = Document()

    # Set page size and margins
    section = doc.sections[0]
    section.page_width = Mm(config.page_width_mm)
    section.page_height = Mm(config.page_height_mm)
    section.top_margin = Mm(config.margin_top_mm)
    section.bottom_margin = Mm(config.margin_bottom_mm)
    section.left_margin = Mm(config.margin_left_mm)
    section.right_margin = Mm(config.margin_right_mm)

    for page_num in range(generator.total_pages):
        if page_num > 0:
            # Add page break
            doc.add_page_break()

        # Create table for the grid
        table = doc.add_table(
            rows=config.labels_per_column,
            cols=config.labels_per_row
        )
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set table properties for no spacing
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)

        grid = generator.get_labels_grid(page_num)

        for row_idx, row in enumerate(grid):
            table_row = table.rows[row_idx]
            table_row.height = Mm(config.label_height_mm)

            for col_idx, label in enumerate(row):
                cell = table_row.cells[col_idx]
                cell.width = Mm(config.label_width_mm)

                # Clear default paragraph
                cell.paragraphs[0].clear()

                if label and not label.is_empty():
                    # Location line 1
                    p1 = cell.paragraphs[0]
                    run1 = p1.add_run(label.location_line1)
                    run1.font.size = Pt(config.font_size_pt)
                    run1.font.name = config.font_family

                    # Location line 2
                    p2 = cell.add_paragraph()
                    run2 = p2.add_run(label.location_line2)
                    run2.font.size = Pt(config.font_size_pt)
                    run2.font.name = config.font_family

                    # Empty line
                    p3 = cell.add_paragraph()
                    p3.add_run("")

                    # Code
                    p4 = cell.add_paragraph()
                    run4 = p4.add_run(label.code)
                    run4.font.size = Pt(config.font_size_pt)
                    run4.font.name = config.font_family

                    # Date
                    p5 = cell.add_paragraph()
                    run5 = p5.add_run(label.date)
                    run5.font.size = Pt(config.font_size_pt)
                    run5.font.name = config.font_family

                    # Additional info (if any)
                    if label.additional_info:
                        p6 = cell.add_paragraph()
                        run6 = p6.add_run(label.additional_info)
                        run6.font.size = Pt(config.font_size_pt * 0.9)
                        run6.font.name = config.font_family
                        run6.italic = True

                    # Set paragraph spacing
                    for p in cell.paragraphs:
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.line_spacing = config.line_spacing

        # Set cell borders
        _set_table_borders(table)

    doc.save(path)

    if open_after:
        webbrowser.open(f"file://{path.absolute()}")

    return path


def _set_table_borders(table):
    """Set light borders on all table cells."""
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return

    tbl = table._tbl
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 0.5pt
        border.set(qn('w:color'), 'CCCCCC')
        tblBorders.append(border)

    tbl.tblPr.append(tblBorders)
