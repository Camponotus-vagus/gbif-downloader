"""
Input handlers for various file formats.

Supports: Excel (.xlsx, .xls), CSV, TXT, DOCX, JSON, YAML
"""

from pathlib import Path
from typing import List, Union
import json

from ..label_generator import Label


def load_data(file_path: Union[str, Path]) -> List[Label]:
    """Load label data from a file.

    Automatically detects the file format based on extension and uses
    the appropriate handler.

    Args:
        file_path: Path to the input file

    Returns:
        List of Label objects

    Raises:
        ValueError: If the file format is not supported
        FileNotFoundError: If the file does not exist
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    extension = path.suffix.lower()

    handlers = {
        ".xlsx": load_excel,
        ".xls": load_excel,
        ".csv": load_csv,
        ".txt": load_txt,
        ".docx": load_docx,
        ".json": load_json,
        ".yaml": load_yaml,
        ".yml": load_yaml,
    }

    handler = handlers.get(extension)
    if handler is None:
        supported = ", ".join(handlers.keys())
        raise ValueError(
            f"Unsupported file format: {extension}. "
            f"Supported formats: {supported}"
        )

    return handler(path)


def load_excel(file_path: Path) -> List[Label]:
    """Load labels from an Excel file (.xlsx, .xls).

    Expected columns (case-insensitive, flexible naming):
    - location_line1 / location1 / località1 / location
    - location_line2 / location2 / località2
    - code / specimen_code / codice
    - date / collection_date / data
    - additional_info / notes / note (optional)
    - count / quantity / quantità (optional, for duplicating labels)
    """
    try:
        import pandas as pd
    except ImportError:
        raise ImportError(
            "pandas and openpyxl are required for Excel support. "
            "Install with: pip install pandas openpyxl"
        )

    df = pd.read_excel(file_path)
    return _dataframe_to_labels(df)


def load_csv(file_path: Path) -> List[Label]:
    """Load labels from a CSV file.

    Same column expectations as Excel files.
    Supports both comma and semicolon delimiters.
    """
    try:
        import pandas as pd
    except ImportError:
        raise ImportError(
            "pandas is required for CSV support. "
            "Install with: pip install pandas"
        )

    # Try comma first, then semicolon
    try:
        df = pd.read_csv(file_path, delimiter=",")
        if len(df.columns) == 1:
            df = pd.read_csv(file_path, delimiter=";")
    except Exception:
        df = pd.read_csv(file_path, delimiter=";")

    return _dataframe_to_labels(df)


def load_txt(file_path: Path) -> List[Label]:
    """Load labels from a TXT file.

    Supports multiple formats:
    1. Tab-separated values (TSV)
    2. Key-value pairs per label (separated by blank lines)
    3. Simple format: each group of 4-5 lines is a label

    Format 2 example:
        location1: Italia, Trentino Alto Adige,
        location2: Giustino (TN), Vedretta d'Amola
        code: N1
        date: 15.vi.2024

        location1: Italia, Trentino Alto Adige,
        location2: Giustino (TN), Vedretta d'Amola
        code: N2
        date: 15.vi.2024
    """
    content = file_path.read_text(encoding="utf-8")
    lines = content.strip().split("\n")

    # Detect format
    if "\t" in lines[0] and not ":" in lines[0]:
        # TSV format
        try:
            import pandas as pd
            df = pd.read_csv(file_path, delimiter="\t")
            return _dataframe_to_labels(df)
        except Exception:
            pass

    # Check for key-value format
    if ":" in content:
        return _parse_key_value_txt(content)

    # Simple line-based format
    return _parse_simple_txt(lines)


def _parse_key_value_txt(content: str) -> List[Label]:
    """Parse key-value formatted text."""
    labels = []
    blocks = content.strip().split("\n\n")

    for block in blocks:
        if not block.strip():
            continue

        data = {}
        for line in block.strip().split("\n"):
            if ":" in line:
                key, value = line.split(":", 1)
                data[key.strip().lower()] = value.strip()

        if data:
            label = Label(
                location_line1=data.get("location1", data.get("location_line1", data.get("località1", ""))),
                location_line2=data.get("location2", data.get("location_line2", data.get("località2", ""))),
                code=data.get("code", data.get("codice", data.get("specimen_code", ""))),
                date=data.get("date", data.get("data", data.get("collection_date", ""))),
                additional_info=data.get("additional_info", data.get("notes", data.get("note", ""))),
            )

            # Handle count/quantity for duplicates
            count = int(data.get("count", data.get("quantity", data.get("quantità", 1))))
            labels.extend([Label(
                location_line1=label.location_line1,
                location_line2=label.location_line2,
                code=label.code,
                date=label.date,
                additional_info=label.additional_info,
            ) for _ in range(count)])

    return labels


def _parse_simple_txt(lines: List[str]) -> List[Label]:
    """Parse simple line-based format (4-5 lines per label)."""
    labels = []
    current_lines = []

    for line in lines:
        line = line.strip()
        if not line:
            if current_lines:
                label = _lines_to_label(current_lines)
                if not label.is_empty():
                    labels.append(label)
                current_lines = []
        else:
            current_lines.append(line)

    # Don't forget the last label
    if current_lines:
        label = _lines_to_label(current_lines)
        if not label.is_empty():
            labels.append(label)

    return labels


def _lines_to_label(lines: List[str]) -> Label:
    """Convert a list of lines to a Label."""
    return Label(
        location_line1=lines[0] if len(lines) > 0 else "",
        location_line2=lines[1] if len(lines) > 1 else "",
        code=lines[2] if len(lines) > 2 else "",
        date=lines[3] if len(lines) > 3 else "",
        additional_info=lines[4] if len(lines) > 4 else "",
    )


def load_docx(file_path: Path) -> List[Label]:
    """Load labels from a Word document (.docx).

    Supports two formats:
    1. Table format: Each row is a label with columns for each field
    2. Paragraph format: Labels separated by blank paragraphs
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX support. "
            "Install with: pip install python-docx"
        )

    doc = Document(file_path)
    labels = []

    # Try table format first
    for table in doc.tables:
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]

        for row in table.rows[1:]:
            data = {headers[i]: cell.text.strip() for i, cell in enumerate(row.cells) if i < len(headers)}
            label = Label.from_dict(data)
            if not label.is_empty():
                labels.append(label)

    # If no tables, try paragraph format
    if not labels:
        paragraphs = [p.text.strip() for p in doc.paragraphs]
        current_lines = []

        for para in paragraphs:
            if not para:
                if current_lines:
                    label = _lines_to_label(current_lines)
                    if not label.is_empty():
                        labels.append(label)
                    current_lines = []
            else:
                current_lines.append(para)

        if current_lines:
            label = _lines_to_label(current_lines)
            if not label.is_empty():
                labels.append(label)

    return labels


def load_json(file_path: Path) -> List[Label]:
    """Load labels from a JSON file.

    Expected format:
    {
        "labels": [
            {
                "location_line1": "Italia, Trentino Alto Adige,",
                "location_line2": "Giustino (TN), Vedretta d'Amola",
                "code": "N1",
                "date": "15.vi.2024",
                "count": 5  // optional, creates duplicates
            },
            ...
        ]
    }

    Or simply an array:
    [
        {"location_line1": "...", ...},
        ...
    ]
    """
    content = file_path.read_text(encoding="utf-8")
    data = json.loads(content)

    if isinstance(data, dict):
        items = data.get("labels", data.get("data", []))
    else:
        items = data

    labels = []
    for item in items:
        label = Label.from_dict(item)
        count = int(item.get("count", item.get("quantity", 1)))
        labels.extend([Label(
            location_line1=label.location_line1,
            location_line2=label.location_line2,
            code=label.code,
            date=label.date,
            additional_info=label.additional_info,
        ) for _ in range(count)])

    return labels


def load_yaml(file_path: Path) -> List[Label]:
    """Load labels from a YAML file.

    Same structure as JSON format.
    """
    try:
        import yaml
    except ImportError:
        raise ImportError(
            "PyYAML is required for YAML support. "
            "Install with: pip install pyyaml"
        )

    content = file_path.read_text(encoding="utf-8")
    data = yaml.safe_load(content)

    if isinstance(data, dict):
        items = data.get("labels", data.get("data", []))
    else:
        items = data

    labels = []
    for item in items:
        label = Label.from_dict(item)
        count = int(item.get("count", item.get("quantity", 1)))
        labels.extend([Label(
            location_line1=label.location_line1,
            location_line2=label.location_line2,
            code=label.code,
            date=label.date,
            additional_info=label.additional_info,
        ) for _ in range(count)])

    return labels


def _dataframe_to_labels(df) -> List[Label]:
    """Convert a pandas DataFrame to a list of Labels."""
    # Normalize column names
    df.columns = [str(c).strip().lower() for c in df.columns]

    # Map various column name variations
    column_map = {
        "location_line1": ["location_line1", "location1", "località1", "location", "loc1"],
        "location_line2": ["location_line2", "location2", "località2", "loc2"],
        "code": ["code", "specimen_code", "codice", "id", "specimen_id"],
        "date": ["date", "collection_date", "data", "data_raccolta"],
        "additional_info": ["additional_info", "notes", "note", "info"],
        "count": ["count", "quantity", "quantità", "n", "copies"],
    }

    def find_column(possible_names):
        for name in possible_names:
            if name in df.columns:
                return name
        return None

    labels = []
    for _, row in df.iterrows():
        data = {}
        for field, possibilities in column_map.items():
            col = find_column(possibilities)
            if col and col in row:
                value = row[col]
                # Handle NaN values
                if hasattr(value, "__float__") and str(value) == "nan":
                    value = ""
                data[field] = str(value) if value else ""

        label = Label(
            location_line1=data.get("location_line1", ""),
            location_line2=data.get("location_line2", ""),
            code=data.get("code", ""),
            date=data.get("date", ""),
            additional_info=data.get("additional_info", ""),
        )

        if not label.is_empty():
            count = 1
            if "count" in data and data["count"]:
                try:
                    count = int(float(data["count"]))
                except (ValueError, TypeError):
                    count = 1

            labels.extend([Label(
                location_line1=label.location_line1,
                location_line2=label.location_line2,
                code=label.code,
                date=label.date,
                additional_info=label.additional_info,
            ) for _ in range(count)])

    return labels
